from firebase import firebase
import threading
import time
from docx import Document
from docx.shared import Inches
from wakeonlan import send_magic_packet
import os
import webbrowser
import subprocess

#Name:Ashutosh Upadhye,Aadesh Keremane
#Python Client script for user computers
def configure_firebase(firebase):
    computer_name = os.getlogin() # Get the username of the computer
    print(computer_name)
    try:
        firebase = firebase.FirebaseApplication('https://echoservice-b8211.firebaseio.com/', None) #Connect to the firebase URL
    except:
        print("Connection has been aborted.") #Connection exception
    data = firebase.get('/names', None) # Get the names in the firebase database
    if not data : print("DataBase Empty")# If no data is found, print empty database
    if data:        #If data is found.....
        for key in data:# Get all the data and load all the records.....
            response = data[key]
            print(response)
            print(key)
            str_data = str(response['name']) # Get the latest records.....
            #String parsing starts.......
            if str_data.__contains__("file") and str_data.__contains__("create") and str_data.__contains__(computer_name.lower()):#If the input commands contains words file and create, then create a file and delete the database record.

                create_file()
                firebase.delete('/names/', key)
               # print('deleted')
                print("File Created.")

            if str_data.__contains__("file") and str_data.__contains__("type") and str_data.__contains__(computer_name.lower()):#If the input commands contains words type and create and file, then create a file and delete the database record and type the daata into the file.

                para = response['name'].replace('type', '')
                stripped_paragraph = str(para).replace('in a file', '').replace('computer', '').replace(computer_name.lower(), '').replace('i want to', '').replace(para[len(para)-1], '')
                create_file(stripped_paragraph)
                firebase.delete('/names/', key)
                #print('deleted')
                print("File Created with paragraph.")

            if str_data.__contains__("shutdown") and str_data.__contains__("computer") and str_data.__contains__(computer_name.lower()):#If the input commands contains words shutdown and computer, shutdown the computer.

                print("shutting down pc")
                firebase.delete('/names/', key)
                #print('deleted')
                shut_down()

            if str_data.__contains__("restart") and str_data.__contains__("computer") and str_data.__contains__(computer_name.lower()):#If the input commands contains words restart and computer, restart the computer.

                print("restarting down pc")
                firebase.delete('/names/', key)
                restart()

            # If the input commands contains words chrome and computer, find the registry key from the regedit of the excel app and find the application path in the users computer .
            # Open the computer using subprocess
            if str_data.__contains__("chrome") and str_data.__contains__("computer") and str_data.__contains__(
                computer_name.lower()) and str_data.__contains__("youtube"):
                print("restarting down pc")
                firebase.delete('/names/', key)
                chrome_path = "C:\\Program Files (x86)\\Google\\Chrome\\Application\\chrome.exe"
                webbrowser.register('chrome', None, webbrowser.BackgroundBrowser(chrome_path), 1)
                webbrowser.get('chrome').open('https://www.youtube.com')

            # If the input commands contains words excel and computer, find the registry key from the regedit of the excel app and find the application path in the users computer .
            #Open the computer using subprocess
            if str_data.__contains__("excel") and str_data.__contains__("computer") and str_data.__contains__(
                    computer_name.lower()) and str_data.__contains__("open"):
                firebase.delete('/names/', key)
                try:
                    import winreg
                except ImportError:
                    import _winreg as winreg

                handle = winreg.OpenKey(winreg.HKEY_LOCAL_MACHINE, r"SOFTWARE\Microsoft\Windows\CurrentVersion\App Paths\excel.exe")

                #num_values = winreg.QueryInfoKey(handle)[1]
                path = winreg.EnumValue(handle, 0)[1]  # Get the value from win registry and evaluate the path required
                subprocess.call(path)  # Use subprocess to call the exe
                # for i in range(num_values):
                #     print(winreg.EnumValue(handle, i))

#Put a firebase database timer ,which reads the database every 5 seconds
def check_firebase_database(n):
    while n > 0:
        time.sleep(2)
        threading.Timer(0, configure_firebase(firebase))
        n = n - 1
        #print("time:"+str(n))


#Create file and insert input statements using docx api into a word document
def create_file(paragraph = None,name = None):
    if paragraph:
        document = Document()
        document.add_paragraph(paragraph, style='Intense Quote')
        document.save('demo.docx')
    if not paragraph:
        document = Document()
        document.save('demo_file.docx')

#Shut down the computer using os.system api and pass a "/s /t " flag to the pc
def shut_down():
    import os
    check = input("Want to shutdown your computer ? (y/n): ");
    if check == 'n':
        exit()
    else:
        os.system("shutdown /s /t 1")
#Restart the computer using os.system api and pass a "/r /t " flag to the pc
def restart():
    import os
    check = input("Want to restart your computer ? (y/n): ");
    if check == 'n':
        exit()
    else:
        os.system("shutdown /r /t 1")


def main():
    #Check the database for 60 seconds...
    check_firebase_database(60)

main()